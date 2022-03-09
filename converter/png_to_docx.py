"""
Provides to insert png file(s) into docx using docx
"""
__author__ = "Aleksandr Shabelsky"
__version__ = "1.3.0"
__email__ = "a.shabelsky@gmail.com"

# Requirement pdf2image: pip install python-docx
# Usage python png_to_docx.py -i 'file1.docx=file1_img_dir'

import os
import time
import docx
import docx.shared
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import argparse
from datetime import datetime
import multiprocessing as mp
from pprint import pprint


def flatten_list(nested_list):
    """
    Function to flatten nested lists
    :param nested_list: list of nested lists
    :return: list
    """
    if not(bool(nested_list)):
        return nested_list
    if isinstance(nested_list[0], list):
        return flatten_list(*nested_list[:1]) + flatten_list(nested_list[1:])
    return nested_list[:1] + flatten_list(nested_list[1:])


def get_img_path(file: str, scan_folders='ENG,RUS_JPG'):
    """
    Function return image folder for file
    :return: str of datetime.now()
    """
    work_dir = os.path.dirname(file)
    img_dir_name = os.path.basename(file).replace("docx", "png").replace('A4(src)', 'A5')
    png_folder = os.path.join(work_dir, scan_folders, img_dir_name)
    if os.path.isdir(png_folder):
        return png_folder


def date_time() -> str:
    """
    Function return string formatted datetime
    :return: str of datetime.now()
    """
    return datetime.now().strftime('%d.%m.%Y %H:%M:%S')


def get_files_list(path: str) -> list:
    """
    Function return list of files in path
    :param: path of files
    :return: list
    """
    return [os.path.join(path, f) for f in os.listdir(path)]


def get_sorted_dict(files: list) -> dict:
    """
    Function return string formatted datetime
    :param: files
    """
    img_dict = {}
    for filename in files:
        if filename.endswith('.png'):
            num_index = filename.rfind('-') + 1
            ext_index = filename.rfind('.png')
            num = filename[num_index:ext_index]
            try:
                n = int(num)
                img_dict[n] = filename
            except Exception as E:
                print(E)
    return img_dict


def png2docx(files: tuple) -> str:
    """
    Function return string formatted datetime
    :param: docx_file
    """
    docx_file, png_folder = files
    doc = docx.Document(docx_file)
    png_files = get_sorted_dict(get_files_list(png_folder))

    picture_width = 148
    picture_height = 210

    picture_num = 0
    paragraph_index = 0

    for paragraph in doc.paragraphs:
        paragraph_index += 1

        if 1 < paragraph_index < 4:
            continue

        if picture_num >= len(png_files):
            if picture_num == 0:
                paragraph.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

            paragraph = paragraph._element
            paragraph.getparent().remove(paragraph)
            paragraph._p = paragraph._element = None

        else:
            paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
            my_run = paragraph.add_run()
            my_run.add_picture(png_files[picture_num + 1],
                               width=docx.shared.Mm(picture_width),
                               height=docx.shared.Mm(picture_height))

        picture_num += 1

    for table in doc.tables:
        paragraph = doc.add_paragraph()
        table._element.addprevious(paragraph._p)
        run = paragraph.add_run()
        run.add_break(docx.enum.text.WD_BREAK.PAGE)

    export_path = os.path.join(os.path.dirname(docx_file), "Edited")

    try:
        os.makedirs(export_path, exist_ok=True)
    except Exception as E:
        print(E)

    docx_rendered_file = os.path.join(export_path, os.path.basename(docx_file).replace('A4(src)', 'A4'))
    doc.save(docx_rendered_file)
    return docx_rendered_file


def main(files, **kwargs):
    """
    Main function
    :param files: docx files to insert png image
    :param kwargs:
    """
    print('#' * 100)
    print(f'[{date_time()}] - Program start')
    start = time.time()

    #  make list of files is flat
    files2convert = [(file, get_img_path(file)) for file in flatten_list(files)
                     if file.endswith('.docx') and os.path.exists(file)]

    if files2convert:
        print(f'[{date_time()}] - Insert PNG to DOCX')

        #  make list of tuples [(file, img_path)]
        pprint(files2convert)

        multiprocessing = kwargs.get('multiprocessing', False)
        if multiprocessing:
            with mp.Pool() as pool:
                print(f'[{date_time()}] - Converting file with multi processing')
                converted_files = pool.map(png2docx, files2convert)
        else:
            print(f'[{date_time()}] - Converting file with single processing')
            converted_files = [png2docx(files=files) for files in files2convert]

        for file in converted_files:
            print(f'[{date_time()}] - Converted file - {file}')
    else:
        print(f'[{date_time()}] - There are no files to converting')
        pprint(files)

    end_time = int(time.time() - start)
    print(f"[{date_time()}] - Time elapsed {end_time} seconds")
    print(f'[{date_time()}] - Program end')
    print('#' * 100)


if __name__ == '__main__':
    # command line argument parser with help message
    desc_msg = "png2docx"
    help_msg = "DOCX files to insert PNG files in self"
    arg_parser = argparse.ArgumentParser(description=desc_msg, formatter_class=argparse.RawTextHelpFormatter)
    arg_parser.add_argument("-i", dest="input", required=True, nargs='+', action='append', help=help_msg)
    arg_parser.add_argument("-m", dest="multiprocessing", required=False, type=bool, default=False)
    args = arg_parser.parse_args()
    print(args)
    main(args.input, multiprocessing=args.multiprocessing)
